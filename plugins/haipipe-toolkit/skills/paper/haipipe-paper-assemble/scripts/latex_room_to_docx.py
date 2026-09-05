"""Build a complete paper from a configured LaTeX desk room.

This is the initial reusable ``latex-room`` adapter for
``haipipe-paper-assemble``. The paper-specific configuration is supplied by
the ``HAIPIPE_PAPER_BUILD_CONFIG`` environment variable, so the same engine
can serve multiple paper rooms. No DOCX is used as an input.
"""

from __future__ import annotations

from dataclasses import dataclass
import hashlib
import json
from pathlib import Path
import re
import shutil
import subprocess
import os
import tomllib
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONFIG_ENV = os.environ.get("HAIPIPE_PAPER_BUILD_CONFIG")
CONFIG_PATH = (
    Path(CONFIG_ENV).expanduser().resolve()
    if CONFIG_ENV
    else Path(__file__).resolve().with_name("paper-build.toml")
)
WORD_ROOM = CONFIG_PATH.parent
ROOT = WORD_ROOM.parent
SKILL_ROOT = Path(__file__).resolve().parent.parent
PROFILE_ROOT = SKILL_ROOT / "profiles"


def load_config() -> dict[str, object]:
    if not CONFIG_PATH.exists():
        return {}
    return tomllib.loads(CONFIG_PATH.read_text(encoding="utf-8"))


BUILD_CONFIG = load_config()
PAPER_CONFIG = BUILD_CONFIG.get("paper", {})
PROFILE_NAME = str(PAPER_CONFIG.get("venue_profile", "")) if isinstance(PAPER_CONFIG, dict) else ""


def load_shared_profile(name: str) -> dict[str, object]:
    """Load venue behaviour from the skill, never manuscript content from a paper."""
    if not name:
        return {}
    path = PROFILE_ROOT / f"{name}.toml"
    if not path.exists():
        raise FileNotFoundError(f"Venue profile not found: {path}")
    return tomllib.loads(path.read_text(encoding="utf-8"))


SHARED_PROFILE = load_shared_profile(PROFILE_NAME)
PAPER_PROFILE = BUILD_CONFIG.get("profile", {})
if not isinstance(PAPER_PROFILE, dict):
    raise TypeError("[profile] must be a TOML table")
PROFILE_CONFIG = {**SHARED_PROFILE, **PAPER_PROFILE}
SOURCE_CONFIG = BUILD_CONFIG.get("source", {})
OUTPUT_CONFIG = BUILD_CONFIG.get("outputs", {})
EVIDENCE_CONFIG = BUILD_CONFIG.get("evidence", {})
if not isinstance(SOURCE_CONFIG, dict) or not isinstance(OUTPUT_CONFIG, dict) or not isinstance(EVIDENCE_CONFIG, dict):
    raise TypeError("[source], [outputs], and [evidence] must be TOML tables")
LATEX_ROOM = (WORD_ROOM / Path(str(SOURCE_CONFIG.get("room", ".")))).resolve()
SECTION_DIR = LATEX_ROOM / str(SOURCE_CONFIG.get("sections", "sections"))
DISPLAY_DIR = LATEX_ROOM / str(SOURCE_CONFIG.get("displays", "displays"))
MASTER = LATEX_ROOM / str(SOURCE_CONFIG.get("master", "main.tex"))
BIB_PATH = LATEX_ROOM / str(SOURCE_CONFIG.get("bibliography", "reference.bib"))


def output_path(key: str, default: str) -> Path:
    value = Path(str(OUTPUT_CONFIG.get(key, default)))
    return value if value.is_absolute() else WORD_ROOM / value


MAIN_PATH = output_path("main_docx", "manuscript-submission-draft.docx")
SUPP_PATH = output_path("supplement_docx", "manuscript-online-supplement-draft.docx")
MAIN_PDF_PATH = output_path("main_pdf", MAIN_PATH.with_suffix(".pdf").name)
SUPP_PDF_PATH = output_path("supplement_pdf", SUPP_PATH.with_suffix(".pdf").name)
DRAFT_SECTION_ROOM = output_path("section_snapshots", "draft-sections")
ASSET_DIR = output_path("assets", "submission-assets")
MANIFEST_PATH = output_path("manifest", "build-manifest.json")
QA_REPORT_PATH = output_path("qa_report", "build-qa.json")
RUNNING_TITLE_FALLBACK = str(
    PROFILE_CONFIG.get(
        "running_title_fallback",
        "",
    )
)

FONT = str(PROFILE_CONFIG.get("font", "Arial"))
CITATION_NUMBERS: dict[str, int] = {}
REF_NUMBERS: dict[str, str] = {}
REF_TEXT: dict[str, str] = {}
BIB: dict[str, dict[str, str]] = {}


@dataclass
class Display:
    kind: str
    block: str
    caption: str = ""
    label: str = ""
    rows: list[list[str]] | None = None
    image_ref: str = ""
    image_path: Path | None = None


@dataclass
class Event:
    kind: str
    value: str | list[str] | Display
    level: int = 0


def parse_group_at(text: str, start: int) -> tuple[str, int] | None:
    """Return the balanced braced group at or after ``start``."""
    pos = start
    while pos < len(text) and text[pos].isspace():
        pos += 1
    if pos >= len(text) or text[pos] != "{":
        return None
    depth = 0
    for index in range(pos, len(text)):
        escaped = index > 0 and text[index - 1] == "\\"
        if text[index] == "{" and not escaped:
            depth += 1
        elif text[index] == "}" and not escaped:
            depth -= 1
            if depth == 0:
                return text[pos + 1:index], index + 1
    return text[pos + 1:], len(text)


def extract_command(text: str, command: str, occurrence: int = 1) -> str:
    pattern = re.compile(r"\\" + re.escape(command) + r"\*?")
    match = None
    cursor = 0
    for _ in range(occurrence):
        match = pattern.search(text, cursor)
        if match is None:
            return ""
        cursor = match.end()
    brace = text.find("{", match.end())
    group = parse_group_at(text, brace)
    return group[0] if group else ""


def unwrap_command(text: str, command: str, arg_count: int, keep: int = -1) -> str:
    """Replace a command with one of its balanced arguments."""
    pattern = re.compile(r"\\" + re.escape(command) + r"\*?")
    cursor = 0
    while True:
        match = pattern.search(text, cursor)
        if match is None:
            return text
        pos = match.end()
        args: list[str] = []
        end = pos
        for _ in range(arg_count):
            group = parse_group_at(text, end)
            if group is None:
                args = []
                break
            value, end = group
            args.append(value)
        if not args:
            cursor = match.end()
            continue
        replacement = args[keep]
        text = text[:match.start()] + replacement + text[end:]
        cursor = match.start() + len(replacement)


def strip_comments(text: str) -> str:
    return re.sub(r"(?m)(?<!\\)%[^\n]*", "", text)


def resolve_input(reference: str, current_base: Path) -> Path | None:
    raw = Path(reference.strip())
    names = [raw] if raw.suffix else [Path(str(raw) + ".tex"), raw]
    candidates: list[Path] = []
    for base in (current_base, LATEX_ROOM, ROOT):
        candidates.extend(base / name for name in names)
    for path in candidates:
        if path.is_file():
            return path.resolve()
    return None


def expand_inputs(text: str, current_base: Path, seen: set[Path] | None = None) -> str:
    seen = set() if seen is None else seen
    pattern = re.compile(r"\\input\s*\{([^}]+)\}")

    def replace(match: re.Match[str]) -> str:
        path = resolve_input(match.group(1), current_base)
        if path is None or path in seen:
            return ""
        seen.add(path)
        return "\n" + expand_inputs(path.read_text(encoding="utf-8"), path.parent, seen) + "\n"

    previous = None
    while previous != text:
        previous = text
        text = pattern.sub(replace, text)
    return text


def parse_bib(text: str) -> dict[str, dict[str, str]]:
    records: dict[str, dict[str, str]] = {}
    for match in re.finditer(r"@\w+\s*\{\s*([^,]+),", text):
        key = match.group(1).strip()
        start = match.end()
        depth = 1
        index = start
        in_quote = False
        while index < len(text) and depth:
            char = text[index]
            if char == '"' and (index == 0 or text[index - 1] != "\\"):
                in_quote = not in_quote
            elif not in_quote:
                if char == "{":
                    depth += 1
                elif char == "}":
                    depth -= 1
            index += 1
        body = text[start:index - 1]
        fields: dict[str, str] = {}
        cursor = 0
        while cursor < len(body):
            field = re.search(r"(\w+)\s*=\s*", body[cursor:])
            if field is None:
                break
            name = field.group(1).lower()
            value_start = cursor + field.end()
            if value_start >= len(body):
                break
            if body[value_start] == "{":
                group = parse_group_at(body, value_start)
                if group is None:
                    break
                value, cursor = group
            elif body[value_start] == '"':
                value_end = value_start + 1
                while value_end < len(body):
                    if body[value_end] == '"' and body[value_end - 1] != "\\":
                        break
                    value_end += 1
                value = body[value_start + 1:value_end]
                cursor = value_end + 1
            else:
                value_end = body.find(",", value_start)
                if value_end < 0:
                    value_end = len(body)
                value = body[value_start:value_end]
                cursor = value_end
            fields[name] = value.strip()
            comma = body.find(",", cursor)
            cursor = len(body) if comma < 0 else comma + 1
        records[key] = fields
    return records


def reference_value(label: str, full: bool = False) -> str:
    if full and label in REF_TEXT:
        return REF_TEXT[label]
    if not full and label in REF_NUMBERS:
        return REF_NUMBERS[label]
    return REF_TEXT.get(label, label.split(":")[-1])


def citation_text(match: re.Match[str]) -> str:
    keys = [key.strip() for key in match.group(1).split(",")]
    numbers: list[str] = []
    for key in keys:
        if key not in BIB:
            continue
        if key not in CITATION_NUMBERS:
            CITATION_NUMBERS[key] = len(CITATION_NUMBERS) + 1
        numbers.append(str(CITATION_NUMBERS[key]))
    return "[" + ",".join(numbers) + "]" if numbers else ""


def latex_to_text(value: str) -> str:
    value = value or ""
    value = re.sub(r"\\cite(?:p|t)?(?:\[[^\]]*\])?\{([^}]+)\}", citation_text, value)
    value = re.sub(
        r"\\(?:cref|Cref)\{([^}]+)\}",
        lambda m: ", ".join(reference_value(label.strip(), full=True) for label in m.group(1).split(",")),
        value,
    )
    value = re.sub(r"\\ref\{([^}]+)\}", lambda m: reference_value(m.group(1)), value)
    value = re.sub(r"\\nameref\{([^}]+)\}", lambda m: reference_value(m.group(1), full=True), value)

    value = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", r"\1/\2", value)
    value = re.sub(r"\\textsuperscript\{2\}", "²", value)
    for command, count, keep in (
        ("textcolor", 2, 1),
        ("href", 2, 1),
        ("multicolumn", 3, 2),
        ("multirow", 3, 2),
        ("resizebox", 3, 2),
        ("makecell", 1, 0),
    ):
        value = unwrap_command(value, command, count, keep)
    for command in (
        "textbf", "textit", "emph", "underline", "textrm", "textsf", "texttt",
        "textsc", "mbox", "mathrm", "mathbf", "operatorname", "text", "intertext",
        "textsuperscript", "textsubscript", "footnote", "thanks",
    ):
        value = unwrap_command(value, command, 1, 0)

    value = re.sub(r"\\textsubscript\{([^{}]*)\}", r"_\1", value)
    value = re.sub(r"\\url\{([^}]*)\}", r"\1", value)
    value = re.sub(r"\\todo\{([^{}]*)\}", r"[TODO: \1]", value)
    value = value.replace(r"\nvzq", "[not independently verified]")
    value = re.sub(r"\\label\{[^}]*\}", "", value)
    value = re.sub(r"\\(?:small|footnotesize|scriptsize|normalsize|large|Large|maketitle|clearpage|newpage|noindent|centering|raggedright)\b", "", value)

    # Convert the small set of TeX accent and annotation forms that occur in
    # the manuscript before the generic command stripper runs.  These are
    # formatting tokens, not alternate manuscript text.
    accent_map = {
        "a": "á", "e": "é", "i": "í", "o": "ó", "u": "ú",
        "A": "Á", "E": "É", "I": "Í", "O": "Ó", "U": "Ú",
        "n": "ñ", "N": "Ñ", "c": "ç", "C": "Ç",
    }
    for letter, accented in accent_map.items():
        value = value.replace(r"\'{" + letter + "}", accented)
        value = value.replace(r"\'" + letter, accented)

    value = re.sub(r"\\[\"`~^]([A-Za-z])", lambda m: m.group(1), value)
    value = re.sub(r"\\c\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\^\s*\{([^{}]*)\}", r"\1", value)

    replacements = {
        "~": " ", r"\&": "&", r"\%": "%", r"\_": "_", r"\#": "#",
        r"\ ": " ", r"\,": " ", r"\;": " ", r"\!": " ", r"\ldots": "...", r"\dots": "...",
        r"\beta": "β", r"\alpha": "α", r"\gamma": "γ", r"\lambda": "λ", r"\rho": "ρ",
        r"\mu": "μ", r"\sigma": "σ", r"\geq": "≥", r"\leq": "≤", r"\approx": "≈",
        r"\times": "×", r"\pm": "±", r"\degree": "°", r"\textendash": "–",
        r"\textemdash": "—", r"\ast": "*", r"\dagger": "†", r"\ddagger": "‡",
        r"\star": "*", "``": '"', "''": '"', "---": "—", "--": "–",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    value = re.sub(r"\\[A-Za-z]+\*?(?:\[[^\]]*\])?", "", value)
    value = value.replace("{", "").replace("}", "").replace("$", "")
    value = value.replace("^2", "²").replace("^3", "³")
    value = value.replace("kg/m2", "kg/m²")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def register_labels(text: str) -> None:
    for match in re.finditer(
        r"\\(?:section|subsection|subsubsection)\*?\s*\{([^{}]+)\}\s*\\label\{([^}]+)\}", text
    ):
        REF_TEXT[match.group(2)] = latex_to_text(match.group(1))

    counters = {"table": 0, "figure": 0}
    pattern = re.compile(r"\\begin\{(figure\*?|table\*?)\}.*?\\end\{\1\}", re.S)
    for match in pattern.finditer(text):
        kind = "figure" if match.group(1).startswith("figure") else "table"
        label = extract_command(match.group(0), "label")
        if not label:
            continue
        counters[kind] += 1
        REF_NUMBERS[label] = str(counters[kind])
        REF_TEXT[label] = f"{kind.title()} {counters[kind]}"


def find_asset(reference: str) -> Path | None:
    ref = Path(reference.strip())
    candidates: list[Path] = []
    if ref.suffix:
        candidates.extend([LATEX_ROOM / ref, LATEX_ROOM / "displays" / ref])
        candidates.extend([p.with_suffix(".png") for p in (LATEX_ROOM / ref, LATEX_ROOM / "displays" / ref)])
    else:
        for ext in (".png", ".jpg", ".jpeg", ".pdf"):
            candidates.extend([LATEX_ROOM / (str(ref) + ext), LATEX_ROOM / "displays" / (str(ref) + ext)])
    for path in candidates:
        if path.is_file():
            return path.resolve()
    stem = ref.stem
    matches = [
        path for path in LATEX_ROOM.rglob("*")
        if path.is_file() and path.stem == stem and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".pdf"}
    ]
    matches.sort(key=lambda path: ({".png": 0, ".jpg": 1, ".jpeg": 2, ".pdf": 3}.get(path.suffix.lower(), 9), str(path)))
    return matches[0].resolve() if matches else None


def parse_table_rows(block: str) -> list[list[str]]:
    # `tabularx` has a width argument plus a column specification; ordinary
    # `tabular`/`longtable` have only the latter.  Keeping those declarations
    # outside the captured body prevents column specs such as ``X r X`` from
    # becoming a spurious first table row in Word.
    match = re.search(
        r"\\begin\{tabularx\}\{[^{}]*\}\{[^{}]*\}(.*?)\\end\{tabularx\}", block, re.S
    )
    if match is None:
        match = re.search(
            r"\\begin\{(tabular|longtable)\}\{[^{}]*\}(.*?)\\end\{\1\}", block, re.S
        )
    if match is None:
        return []
    body = match.group(1) if "tabularx" in match.group(0).split("}", 1)[0] else match.group(2)
    body = re.sub(
        r"\\(?:toprule|midrule|bottomrule|hline|addlinespace|cline\{[^}]*\}|cmidrule(?:\([^)]*\))?\{[^}]*\})",
        "",
        body,
    )
    rows: list[list[str]] = []
    for raw in re.split(r"(?<!\\)\\\\(?:\s*\[[^]]*\])?", body):
        raw = raw.strip()
        if not raw or raw.startswith("\\"):
            continue
        cells: list[str] = []
        cursor = 0
        for index, char in enumerate(raw):
            if char == "&" and (index == 0 or raw[index - 1] != "\\"):
                cells.append(latex_to_text(raw[cursor:index]))
                cursor = index + 1
        cells.append(latex_to_text(raw[cursor:]))
        if any(cells):
            rows.append(cells)
    width = max((len(row) for row in rows), default=0)
    return [row + [""] * (width - len(row)) for row in rows]


def parse_display(block: str) -> Display:
    is_figure = bool(re.search(r"\\includegraphics(?:\[[^]]*\])?\{", block))
    caption = extract_command(block, "caption") or extract_command(block, "caption*")
    label = extract_command(block, "label")
    if is_figure:
        image_match = re.search(r"\\includegraphics(?:\[[^]]*\])?\{([^}]+)\}", block)
        image_ref = image_match.group(1) if image_match else ""
        return Display(
            kind="figure",
            block=block,
            caption=latex_to_text(caption),
            label=label,
            image_ref=image_ref,
            image_path=find_asset(image_ref) if image_ref else None,
        )
    return Display(kind="table", block=block, caption=latex_to_text(caption), label=label, rows=parse_table_rows(block))


DISPLAY_PATTERN = re.compile(r"\\begin\{(figure\*?|table\*?)\}.*?\\end\{\1\}", re.S)
LIST_PATTERN = re.compile(r"\\begin\{(itemize|enumerate)\}.*?\\end\{\1\}", re.S)
CENTER_PATTERN = re.compile(r"\\begin\{center\}.*?\\end\{center\}", re.S)
TOKEN_PATTERN = re.compile(r"@@(?:DISPLAY|LIST)\d+@@")


def list_items(block: str) -> list[str]:
    body = re.sub(r"\\begin\{(?:itemize|enumerate)\}|\\end\{(?:itemize|enumerate)\}", "", block)
    return [latex_to_text(match.group(1)) for match in re.finditer(r"\\item\s*(.*?)(?=\\item|$)", body, re.S) if latex_to_text(match.group(1))]


def parse_events(text: str) -> tuple[list[Event], list[Display]]:
    text = re.sub(r"\\(?:clearpage|newpage|pagebreak)\b", "\n", text)
    text = re.sub(r"\\(?:maketitle|thispagestyle|pagestyle)\s*(?:\{[^}]*\})?", "", text)
    text = re.sub(r"\\(?:bibliographystyle|bibliography)\s*(?:\{[^}]*\})?", "", text)
    # Counter and naming declarations control TeX's rendering; they are not
    # manuscript prose and must not leak into a Word projection.
    text = re.sub(r"(?m)^\s*\\(?:setcounter|renewcommand)\b.*$", "", text)
    text = re.sub(r"\\(?:begin|end)\{(?:document|abstract)\}", "", text)
    blocks: dict[str, tuple[str, str]] = {}
    displays: list[Display] = []
    lists: list[list[str]] = []

    def store_display(match: re.Match[str]) -> str:
        token = f"@@DISPLAY{len(blocks)}@@"
        blocks[token] = ("display", match.group(0))
        return "\n" + token + "\n"

    text = DISPLAY_PATTERN.sub(store_display, text)

    def store_center(match: re.Match[str]) -> str:
        block = match.group(0)
        if "tabular" not in block and "includegraphics" not in block:
            return block
        token = f"@@DISPLAY{len(blocks)}@@"
        blocks[token] = ("display", block)
        return "\n" + token + "\n"

    text = CENTER_PATTERN.sub(store_center, text)

    def store_list(match: re.Match[str]) -> str:
        token = f"@@LIST{len(blocks)}@@"
        blocks[token] = ("list", match.group(0))
        return "\n" + token + "\n"

    text = LIST_PATTERN.sub(store_list, text)

    ordered: list[Event] = []
    heading_pattern = re.compile(r"\\(section|subsection|subsubsection|paragraph)\*?\s*")

    def add_ordered_plain(value: str) -> None:
        value = re.sub(r"\\appendix\b", "", value)
        value = re.sub(r"\\(?:begin|end)\{(?:center|table\*?|figure\*?)\}", "", value)
        for paragraph in re.split(r"\n\s*\n+", value):
            paragraph = latex_to_text(paragraph)
            if paragraph:
                ordered.append(Event("text", paragraph))

    def add_before_tokens(value: str) -> None:
        parts = re.split(r"(@@(?:DISPLAY|LIST)\d+@@)", value)
        for part in parts:
            if not part:
                continue
            if part in blocks:
                kind, block = blocks[part]
                if kind == "list":
                    ordered.append(Event("list", list_items(block)))
                else:
                    display = parse_display(block)
                    displays.append(display)
                    ordered.append(Event("display", display))
            else:
                add_ordered_plain(part)

    cursor = 0
    while True:
        match = heading_pattern.search(text, cursor)
        if match is None:
            add_before_tokens(text[cursor:])
            break
        add_before_tokens(text[cursor:match.start()])
        brace = text.find("{", match.end())
        group = parse_group_at(text, brace)
        if group is None:
            cursor = match.end()
            continue
        title, cursor = group
        level = {"section": 1, "subsection": 2, "subsubsection": 3, "paragraph": 3}[match.group(1)]
        ordered.append(Event("heading", latex_to_text(title), level))
    return ordered, displays


def set_font(run, name: str = FONT, size: float = 12, bold: bool | None = None, italic: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document, *, body_size: float = 12, line_spacing: float = 2.0) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    for name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Caption", "List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10 if name == "Caption" else body_size)
        style.paragraph_format.line_spacing = 1.0 if name == "Caption" else line_spacing
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
    for name, size, before in (("Heading 1", body_size, 10), ("Heading 2", body_size, 6), ("Heading 3", body_size, 4)):
        style = doc.styles[name]
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.keep_with_next = True
    doc.styles["Caption"].font.italic = True
    doc.styles["Caption"].font.color.rgb = RGBColor(68, 68, 68)
    doc.styles["Caption"].paragraph_format.space_before = Pt(4)
    doc.styles["Caption"].paragraph_format.space_after = Pt(6)
    doc.styles["List Bullet"].paragraph_format.left_indent = Inches(0.25)
    doc.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.15)
    if bool(PROFILE_CONFIG.get("line_numbers", False)):
        set_line_numbering(section)


def set_line_numbering(section) -> None:
    """Apply continuous Word line numbers when the selected venue requires them."""
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is None:
        existing = OxmlElement("w:lnNumType")
        sect_pr.append(existing)
    existing.set(qn("w:countBy"), "1")
    existing.set(qn("w:distance"), str(PROFILE_CONFIG.get("line_number_distance", 360)))
    existing.set(qn("w:restart"), "newPage" if bool(PROFILE_CONFIG.get("restart_line_numbers_each_page", False)) else "continuous")


def add_text(doc: Document, text: str, *, size: float = 12, bold: bool = False, align=None, style: str = "Normal") -> None:
    if not text:
        return
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    action = "[AUTHOR ACTION REQUIRED" in text or "[TODO:" in text or "[not independently verified]" in text
    set_font(run, size=size, bold=bold, color="C00000" if action else None)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    if not text:
        return
    heading_style = str(PROFILE_CONFIG.get("heading_style", "heading")).lower()
    style = "Normal" if heading_style == "normal" else f"Heading {min(level, 3)}"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.keep_with_next = bool(PROFILE_CONFIG.get("headings_keep_with_next", True))
    paragraph.paragraph_format.line_spacing = 2.0
    # An explicit before-space keeps LibreOffice's DOCX renderer from visually
    # joining an immediately preceding body paragraph and its next heading.
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.upper() if level == 1 else text)
    set_font(run, size=12, bold=True)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 2.0
    run = paragraph.add_run(text)
    set_font(run, size=12)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(index, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc: Document, display: Display, title: str = "", *, compact: bool = False) -> None:
    rows = display.rows or []
    if title or display.caption:
        caption = f"{title}. " if title else ""
        caption += display.caption
        add_text(doc, caption.strip(), size=10 if compact else 11, bold=True, style="Caption")
    if not rows:
        return
    widths = [max(900, 9360 // len(rows[0]))] * len(rows[0])
    widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_font(run, size=8 if compact else 8.5, bold=(row_index == 0))
            if row_index == 0:
                set_cell_shading(cell, "F2F4F7")
        if row_index == 0:
            mark_header_row(table.rows[row_index])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def ensure_asset(display: Display, name: str) -> Path | None:
    if display.image_path is None:
        return None
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    target = ASSET_DIR / name
    source = display.image_path
    if source.suffix.lower() == ".pdf":
        subprocess.run(
            ["pdftoppm", "-png", "-singlefile", "-r", "180", str(source), str(target.with_suffix(""))],
            check=True,
        )
        return target
    shutil.copy2(source, target)
    return target


def add_figure(doc: Document, display: Display, title: str = "", *, compact: bool = False) -> None:
    base_name = title or Path(display.image_ref).stem or "figure"
    filename = re.sub(r"[^A-Za-z0-9._-]+", "-", base_name).strip("-") + ".png"
    image = ensure_asset(display, filename)
    if image is None:
        add_text(doc, f"{title}: image source not found.", size=10 if compact else 11)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(image), width=Inches(6.1))
    alt = display.caption or title or display.image_ref
    shape._inline.docPr.set("descr", alt[:500])
    shape._inline.docPr.set("title", filename)
    caption = f"{title}. " if title else ""
    caption += display.caption
    add_text(doc, caption.strip(), size=10 if compact else 11, style="Caption")


def add_events(doc: Document, events: Iterable[Event], *, include_displays: bool = True, compact: bool = False) -> None:
    for event in events:
        if event.kind == "heading":
            add_heading(doc, str(event.value), event.level)
        elif event.kind == "text":
            add_text(doc, str(event.value), size=11 if compact else 12)
        elif event.kind == "list":
            for item in event.value:  # type: ignore[union-attr]
                add_bullet(doc, str(item))
        elif event.kind == "display" and include_displays:
            display = event.value
            if isinstance(display, Display):
                if display.kind == "table":
                    add_table(doc, display, compact=compact)
                else:
                    add_figure(doc, display, compact=compact)


def add_abstract(doc: Document, block: str) -> None:
    add_heading(doc, str(PROFILE_CONFIG.get("abstract_heading", "ABSTRACT")), 1)
    labels = list(re.finditer(r"\\noindent\s*\\textbf\s*\{([^{}]+)\}", block))
    if not labels:
        add_text(doc, latex_to_text(block))
        return
    for index, match in enumerate(labels):
        end = labels[index + 1].start() if index + 1 < len(labels) else len(block)
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(latex_to_text(match.group(1)) + " ")
        body_run = paragraph.add_run(latex_to_text(block[match.end():end]))
        set_font(label_run, size=12, bold=True)
        set_font(body_run, size=12)


def extract_highlights(body: str) -> list[str]:
    heading = re.escape(str(PROFILE_CONFIG.get("highlights_heading", "Article Highlights")))
    match = re.search(rf"\\textbf\{{{heading}\}}.*?\\begin\{{itemize\}}(.*?)\\end\{{itemize\}}", body, re.S)
    if match is None:
        return []
    return list_items(match.group(0))


def add_title_page(doc: Document, title: str, running_title: str, word_count: int, table_count: int, figure_count: int, author: str) -> None:
    add_text(doc, title, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Running title: {running_title}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, author or "[AUTHOR NAMES — first, middle, and last names]", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, str(PROFILE_CONFIG.get("affiliations_placeholder", "[AFFILIATIONS]")), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, str(PROFILE_CONFIG.get("corresponding_author_placeholder", "Corresponding author: [FULL NAME]; [email]")), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Article type: {PROFILE_CONFIG.get('article_type', 'Research Article')}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Main-text word count (excluding tables, legends, title page, acknowledgments, and references): {word_count}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Number of main displays: {table_count + figure_count} ({table_count} tables, {figure_count} figures)", align=WD_ALIGN_PARAGRAPH.CENTER)
    action_fields = PROFILE_CONFIG.get("title_page_action_fields", [])
    if not isinstance(action_fields, list):
        action_fields = []
    for field in action_fields:
        add_text(doc, str(field), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def labelled_blocks(doc: Document, heading: str, block: str) -> None:
    """Render JAMA Key Points or a structured abstract from labelled TeX prose."""
    add_heading(doc, heading, 1)
    labels = list(re.finditer(r"\\noindent\s*\\textbf\s*\{([^{}]+)\}", block))
    if not labels:
        add_text(doc, latex_to_text(block))
        return
    for index, match in enumerate(labels):
        end = labels[index + 1].start() if index + 1 < len(labels) else len(block)
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(latex_to_text(match.group(1)) + " ")
        body_run = paragraph.add_run(latex_to_text(block[match.end():end]))
        set_font(label_run, size=12, bold=True)
        set_font(body_run, size=12)


def extract_jama_title(center_block: str) -> tuple[str, str]:
    """Read the title and blinded-author line from JAMA's source-owned title block."""
    title_part = center_block.split(r"\vspace", 1)[0]
    title_part = re.sub(r"\\(?:begin|end)\{center\}", "", title_part)
    title_part = re.sub(r"\\(?:Large|large|bfseries)\b", "", title_part)
    title = latex_to_text(title_part.replace(r"\\", " "))
    author_match = re.search(r"\vspace\{[^}]+\}\s*(.*?)\s*\vspace", center_block, re.S)
    author = latex_to_text(author_match.group(1).strip()) if author_match else ""
    return title, author


def jama_parts(raw_master: str) -> tuple[str, str, str, str, str, str]:
    """Split a JAMA IM desk room without requiring a generic abstract environment."""
    body = clean_fragment(raw_master, MASTER.parent, remove_abstract=False)
    body = body.split(r"\begin{document}", 1)[-1].split(r"\end{document}", 1)[0]
    center_match = CENTER_PATTERN.search(body)
    if center_match is None:
        raise RuntimeError("JAMA IM source is missing its title-page center block")
    title, author = extract_jama_title(center_match.group(0))

    key_match = re.search(r"\\section\*\{Key Points\}(.*?)\\section\*\{Abstract\}", body, re.S)
    abstract_match = re.search(r"\\section\*\{Abstract\}(.*?)\\section\s*\{Introduction\}", body, re.S)
    bib_match = re.search(r"\\bibliography\s*\{", body)
    appendix_match = re.search(r"\\appendix\b", body)
    if abstract_match is None or bib_match is None or appendix_match is None:
        raise RuntimeError("JAMA IM source is missing Key Points, Abstract, bibliography, or appendix boundary")
    key_points = key_match.group(1) if key_match else ""
    abstract = abstract_match.group(1)
    intro = re.search(r"\\section\s*\{Introduction\}", body)
    if intro is None:
        raise RuntimeError("JAMA IM source is missing Introduction")
    main_text = body[intro.start():bib_match.start()]
    supplement_text = body[appendix_match.end():]
    return title, author, key_points, abstract, main_text, supplement_text


def add_jama_title_page(doc: Document, title: str, author: str, main_words: int,
                        table_count: int, figure_count: int, evidence: dict[str, object]) -> None:
    """Venue-only title packaging; all study-specific text remains in the room source."""
    add_text(doc, title, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if str(evidence.get("mode", "draft")) not in {"final", "submission", "submission-ready"}:
        add_text(doc, str(PROFILE_CONFIG.get("draft_label", "DRAFT — NOT SUBMISSION READY")),
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    if author:
        add_text(doc, author, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, str(PROFILE_CONFIG.get("venue_label", "JAMA Internal Medicine")),
             align=WD_ALIGN_PARAGRAPH.CENTER)
    date = str(PROFILE_CONFIG.get("document_date", "")).strip()
    if date:
        add_text(doc, date, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Main-text source word count: {main_words}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"Main displays: {table_count} tables; {figure_count} figures", align=WD_ALIGN_PARAGRAPH.CENTER)
    if evidence.get("pending_markers"):
        add_text(doc, "Evidence placeholders remain; this build is not submission-ready.",
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def initials(author: str) -> str:
    tokens = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]+", author)
    return "".join(token[0].upper() for token in tokens if token)


def format_authors(author_field: str) -> str:
    clean = latex_to_text(author_field)
    if not clean:
        return ""
    output: list[str] = []
    for author in re.split(r"\s+and\s+", clean):
        author = author.strip().rstrip(".")
        if author.lower() == "others":
            continue
        if "," in author:
            surname, given = [part.strip() for part in author.split(",", 1)]
            output.append(f"{surname} {initials(given)}".strip())
        else:
            parts = author.split()
            output.append(f"{parts[-1]} {initials(' '.join(parts[:-1]))}".strip() if len(parts) > 1 else author)
    if len(output) > 6:
        return ", ".join(output[:6]) + ", et al"
    return ", ".join(output)


def format_reference(key: str) -> str:
    fields = BIB[key]
    authors = format_authors(fields.get("author", ""))
    title = latex_to_text(fields.get("title", ""))
    journal = latex_to_text(fields.get("journal", fields.get("booktitle", "")))
    year = latex_to_text(fields.get("year", ""))
    volume = latex_to_text(fields.get("volume", ""))
    pages = latex_to_text(fields.get("pages", ""))
    doi = latex_to_text(fields.get("doi", ""))
    pieces = [part for part in (authors, title, journal) if part]
    tail = " ".join(part for part in (year, volume, pages) if part)
    if tail:
        pieces.append(tail)
    if doi:
        pieces.append("doi:" + doi)
    return ". ".join(pieces).rstrip(".") + "."


def add_references(doc: Document) -> None:
    if not CITATION_NUMBERS:
        return
    add_heading(doc, str(PROFILE_CONFIG.get("references_heading", "REFERENCES")), 1)
    for number, key in enumerate(CITATION_NUMBERS, start=1):
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        run = paragraph.add_run(f"{number}. {format_reference(key)}")
        set_font(run, size=11)


def clean_fragment(raw: str, base: Path, *, remove_abstract: bool = True) -> str:
    value = expand_inputs(raw, base)
    value = strip_comments(value)
    if remove_abstract:
        value = re.sub(r"\\(?:begin|end)\{abstract\}", "", value)
    return value


def word_count(events: Iterable[Event]) -> int:
    words = []
    for event in events:
        if event.kind in {"text", "list"}:
            values = event.value if event.kind == "list" else [event.value]
            words.extend(re.findall(r"\b[\w%]+(?:[–-][\w%]+)?\b", " ".join(str(v) for v in values)))
    return len(words)


def write_section_snapshots(section_specs: list[tuple[str, str]], running_title: str) -> None:
    DRAFT_SECTION_ROOM.mkdir(parents=True, exist_ok=True)
    readme = DRAFT_SECTION_ROOM / "README.md"
    readme.write_text(
        "# Generated section snapshots\n\n"
        "These DOCX files are generated from the configured desk-room sections by the shared "
        "`haipipe-paper-assemble` engine, launched through this paper's thin wrapper. "
        "They are review snapshots only and are never used as builder inputs.\n",
        encoding="utf-8",
    )
    for source_name, label in section_specs:
        source = SECTION_DIR / source_name
        text = clean_fragment(source.read_text(encoding="utf-8"), source.parent)
        events, _ = parse_events(text)
        doc = Document()
        configure_document(doc)
        add_title_page(doc, label, running_title, word_count(events), 0, 0, "Authors hidden for review")
        add_events(doc, events, include_displays=True)
        output_name = source_name.replace("_", "-").replace(".tex", ".docx")
        doc.save(DRAFT_SECTION_ROOM / output_name)


def evidence_lock_path() -> Path | None:
    """Return the materialized room-local evidence receipt, when configured."""
    raw = EVIDENCE_CONFIG.get("lock")
    if not raw:
        return None
    path = Path(str(raw))
    return path if path.is_absolute() else (WORD_ROOM / path).resolve()


def evidence_preflight() -> dict[str, object]:
    """Validate evidence state without creating or replacing manuscript prose.

    The Section projection owns the words.  The lock records which Page Evidence
    Items those words depend on.  A draft can keep visible ``[E## pending]``
    markers; a final build cannot.
    """
    mode = str(EVIDENCE_CONFIG.get("mode", "none")).lower()
    path = evidence_lock_path()
    source_text = "\n".join(path.read_text(encoding="utf-8", errors="replace")
                            for path in sorted(SECTION_DIR.glob("*.tex")))
    pending_markers = sorted(set(re.findall(r"\[E\d{2}\s+pending\]", source_text)))
    report: dict[str, object] = {
        "mode": mode,
        "lock": str(path.relative_to(ROOT)) if path and path.is_relative_to(ROOT) else (str(path) if path else None),
        "pending_markers": pending_markers,
        "items": 0,
        "unaccepted_items": [],
        "lock_exists": path.exists() if path else None,
    }
    if path is None:
        return report
    if not path.exists():
        raise FileNotFoundError(f"Configured evidence lock not found: {path}")
    lock = json.loads(path.read_text(encoding="utf-8"))
    items = lock.get("items", [])
    if not isinstance(items, list):
        raise ValueError(f"Evidence lock has non-list items: {path}")
    unaccepted = [
        f"{item.get('page', 'unknown')}/{item.get('item', 'unknown')} ({item.get('state', 'unknown')})"
        for item in items if isinstance(item, dict) and item.get("state") != "accepted"
    ]
    report["items"] = len(items)
    report["unaccepted_items"] = unaccepted
    if mode in {"final", "submission", "submission-ready"} and (pending_markers or unaccepted):
        raise RuntimeError(
            "Final build blocked by unresolved evidence: "
            + ", ".join(pending_markers + unaccepted)
        )
    return report


def source_manifest() -> dict[str, object]:
    paths = [MASTER, BIB_PATH]
    paths.extend(sorted(SECTION_DIR.glob("*.tex")))
    paths.extend(sorted(DISPLAY_DIR.rglob("*.tex")))
    paths.extend(sorted(DISPLAY_DIR.rglob("*.png")))
    paths.extend(sorted(DISPLAY_DIR.rglob("*.pdf")))
    lock_path = evidence_lock_path()
    if lock_path and lock_path.exists():
        paths.append(lock_path)
    files = []
    for path in paths:
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        files.append({"path": str(path.relative_to(ROOT)), "sha256": digest})
    config_hash = hashlib.sha256(CONFIG_PATH.read_bytes()).hexdigest() if CONFIG_PATH.exists() else None
    profile_path = PROFILE_ROOT / f"{PROFILE_NAME}.toml" if PROFILE_NAME else None
    profile_hash = (
        hashlib.sha256(profile_path.read_bytes()).hexdigest()
        if profile_path and profile_path.exists() else None
    )
    output_paths = [MAIN_PATH, SUPP_PATH, MAIN_PDF_PATH, SUPP_PDF_PATH]
    output_hashes = {
        str(path.relative_to(ROOT)): hashlib.sha256(path.read_bytes()).hexdigest()
        for path in output_paths if path.exists()
    }
    return {
        "engine": "haipipe-paper-assemble/latex_room_to_docx-0.2.1",
        "builder": "haipipe-paper-assemble/scripts/latex_room_to_docx.py",
        "config": str(CONFIG_PATH.relative_to(ROOT)) if CONFIG_PATH.exists() else None,
        "config_sha256": config_hash,
        "profile": str(profile_path.relative_to(SKILL_ROOT)) if profile_path and profile_path.exists() else None,
        "profile_sha256": profile_hash,
        "venue_profile": PROFILE_CONFIG.get("venue_profile", BUILD_CONFIG.get("paper", {}).get("venue_profile")),
        "source_of_record": str(MASTER.relative_to(ROOT)),
        "outputs": [
            str(MAIN_PATH.relative_to(ROOT)),
            str(SUPP_PATH.relative_to(ROOT)),
            str(MAIN_PDF_PATH.relative_to(ROOT)),
            str(SUPP_PDF_PATH.relative_to(ROOT)),
            str(DRAFT_SECTION_ROOM.relative_to(ROOT)),
            str(MANIFEST_PATH.relative_to(ROOT)),
            str(QA_REPORT_PATH.relative_to(ROOT)),
        ],
        "output_sha256": output_hashes,
        "files": files,
    }


def render_submission_pdfs() -> dict[str, object]:
    """Render configured DOCX deliverables to PDFs when LibreOffice is available."""
    soffice = shutil.which("soffice")
    report: dict[str, object] = {
        "renderer": "LibreOffice" if soffice else None,
        "available": bool(soffice),
        "outputs": [],
        "errors": [],
    }
    if not soffice:
        return report
    for docx_path, pdf_path in ((MAIN_PATH, MAIN_PDF_PATH), (SUPP_PATH, SUPP_PDF_PATH)):
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            produced = pdf_path.parent / docx_path.with_suffix(".pdf").name
            if produced != pdf_path and produced.exists():
                shutil.move(str(produced), str(pdf_path))
            if not pdf_path.exists():
                raise RuntimeError(f"LibreOffice did not create {pdf_path.name}")
            report["outputs"].append(str(pdf_path.relative_to(ROOT)))  # type: ignore[index]
        except (OSError, subprocess.CalledProcessError, RuntimeError) as error:
            report["errors"].append(f"{docx_path.name}: {error}")  # type: ignore[index]
    return report


def write_common_receipts(main_events: list[Event], main_displays: list[Display],
                          supplement_events: list[Event], running_title: str,
                          evidence: dict[str, object], pdf_report: dict[str, object]) -> None:
    """Write derived receipts only; they are never inputs to a later build."""
    snapshot_config = BUILD_CONFIG.get("snapshots", {})
    configured_sections = snapshot_config.get("sections", []) if isinstance(snapshot_config, dict) else []
    section_specs = [
        (str(item["source"]), str(item["label"]))
        for item in configured_sections
        if isinstance(item, dict) and "source" in item and "label" in item
    ]
    if not section_specs:
        section_specs = [
            (path.name, path.stem.replace("_", " ").title())
            for path in sorted(SECTION_DIR.glob("*.tex"))
        ]
    write_section_snapshots(section_specs, running_title)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    manifest = source_manifest()
    manifest["evidence"] = evidence
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    word_limit = PROFILE_CONFIG.get("main_text_word_limit")
    qa_report = {
        "status": "DRAFT" if evidence.get("pending_markers") or evidence.get("unaccepted_items") else "CANDIDATE",
        "source_of_record": str(MASTER.relative_to(ROOT)),
        "venue_profile": PROFILE_NAME,
        "main_text_words": word_count(main_events),
        "main_text_word_limit": word_limit,
        "citations": len(CITATION_NUMBERS),
        "main_tables": len([display for display in main_displays if display.kind == "table"]),
        "main_figures": len([display for display in main_displays if display.kind == "figure"]),
        "supplement_events": len(supplement_events),
        "evidence": evidence,
        "pdf": pdf_report,
        "checks": {
            "master_exists": MASTER.exists(),
            "bibliography_exists": BIB_PATH.exists(),
            "citations_resolved": True,
            "within_declared_word_limit": (
                word_limit is None or word_count(main_events) <= int(word_limit)
            ),
            "evidence_lock_exists": evidence.get("lock_exists"),
            "no_pending_evidence": not evidence.get("pending_markers") and not evidence.get("unaccepted_items"),
            "pdf_outputs_rendered": not pdf_report.get("available") or not pdf_report.get("errors"),
            "g6_human_decision": False,
        },
    }
    QA_REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    QA_REPORT_PATH.write_text(json.dumps(qa_report, indent=2) + "\n", encoding="utf-8")


def build_jama_internal_medicine(raw_master: str, evidence: dict[str, object]) -> tuple[Path, Path]:
    """JAMA IM is a venue profile, not a paper-specific renderer."""
    title, author, key_points, abstract, main_text, supplement_text = jama_parts(raw_master)
    main_events, main_displays = parse_events(main_text)
    supplement_events, _ = parse_events(supplement_text)
    missing = [key for key in CITATION_NUMBERS if key not in BIB]
    if missing:
        raise RuntimeError(f"Citations missing from reference.bib: {missing}")
    main_tables = [display for display in main_displays if display.kind == "table"]
    main_figures = [display for display in main_displays if display.kind == "figure"]
    MAIN_PATH.parent.mkdir(parents=True, exist_ok=True)
    SUPP_PATH.parent.mkdir(parents=True, exist_ok=True)

    main_doc = Document()
    configure_document(main_doc)
    add_jama_title_page(main_doc, title, author, word_count(main_events), len(main_tables), len(main_figures), evidence)
    if key_points:
        labelled_blocks(main_doc, str(PROFILE_CONFIG.get("key_points_heading", "KEY POINTS")), key_points)
    labelled_blocks(main_doc, str(PROFILE_CONFIG.get("abstract_heading", "ABSTRACT")), abstract)
    add_events(main_doc, main_events, include_displays=bool(PROFILE_CONFIG.get("include_main_displays", False)))
    add_references(main_doc)
    main_doc.save(MAIN_PATH)

    supp_doc = Document()
    configure_document(supp_doc, body_size=11, line_spacing=1.5)
    add_text(supp_doc, title, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(supp_doc, str(PROFILE_CONFIG.get("supplement_title", "ONLINE-ONLY SUPPLEMENTAL MATERIAL")),
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    supp_doc.add_page_break()
    add_events(supp_doc, supplement_events, include_displays=True, compact=True)
    supp_doc.save(SUPP_PATH)

    pdf_report = render_submission_pdfs()
    write_common_receipts(main_events, main_displays, supplement_events, title, evidence, pdf_report)
    return MAIN_PATH, SUPP_PATH


def build() -> tuple[Path, Path]:
    if not MASTER.exists():
        raise FileNotFoundError(MASTER)
    if not BIB_PATH.exists():
        raise FileNotFoundError(BIB_PATH)

    global BIB, CITATION_NUMBERS, REF_NUMBERS, REF_TEXT
    BIB = parse_bib(BIB_PATH.read_text(encoding="utf-8"))
    CITATION_NUMBERS = {}
    REF_NUMBERS = {}
    REF_TEXT = {}

    raw_master = MASTER.read_text(encoding="utf-8")
    evidence = evidence_preflight()
    if str(PROFILE_CONFIG.get("layout", "")).lower() == "jama-internal-medicine":
        return build_jama_internal_medicine(raw_master, evidence)
    running_match = re.search(r"Running title\s*\(<[^>]+>\):\s*(.+)", raw_master)
    running_title = running_match.group(1).strip() if running_match else RUNNING_TITLE_FALLBACK
    title = latex_to_text(extract_command(raw_master, "title"))
    author = latex_to_text(extract_command(raw_master, "author"))
    body = clean_fragment(raw_master, MASTER.parent, remove_abstract=False)
    body = body.split(r"\begin{document}", 1)[-1].split(r"\end{document}", 1)[0]
    register_labels(body)

    abstract_match = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", body, re.S)
    if abstract_match is None:
        raise RuntimeError("Master manuscript has no abstract environment")
    abstract = abstract_match.group(1)
    highlights = extract_highlights(body)
    backmatter_heading = str(PROFILE_CONFIG.get("backmatter_heading", "Acknowledgments"))
    ack_match = re.search(r"\\section\*?\s*\{" + re.escape(backmatter_heading) + r"\}", body)
    bib_match = re.search(r"\\bibliographystyle\s*\{", body)
    appendix_match = re.search(r"\\appendix\b", body)
    if ack_match is None or bib_match is None or appendix_match is None:
        raise RuntimeError("Master manuscript is missing acknowledgments, bibliography, or appendix boundary")

    main_text = body[abstract_match.end():ack_match.start()]
    back_text = body[ack_match.start():bib_match.start()]
    supplement_text = body[appendix_match.end():]
    main_events, main_displays = parse_events(main_text)
    back_events, _ = parse_events(back_text)
    supplement_events, _ = parse_events(supplement_text)
    missing = [key for key in CITATION_NUMBERS if key not in BIB]
    if missing:
        raise RuntimeError(f"Citations missing from reference.bib: {missing}")

    main_tables = [display for display in main_displays if display.kind == "table"]
    main_figures = [display for display in main_displays if display.kind == "figure"]
    MAIN_PATH.parent.mkdir(parents=True, exist_ok=True)
    SUPP_PATH.parent.mkdir(parents=True, exist_ok=True)

    main_doc = Document()
    configure_document(main_doc)
    add_title_page(main_doc, title, running_title, word_count(main_events), len(main_tables), len(main_figures), author)
    add_abstract(main_doc, abstract)
    add_heading(main_doc, "Article Highlights", 1)
    for highlight in highlights:
        add_bullet(main_doc, highlight)
    add_events(main_doc, main_events, include_displays=False)
    add_events(main_doc, back_events, include_displays=False)
    add_references(main_doc)
    add_heading(main_doc, "TABLES", 1)
    for number, display in enumerate(main_tables, start=1):
        add_table(main_doc, display, f"Table {number}")
    add_heading(main_doc, "FIGURE LEGENDS", 1)
    for number, display in enumerate(main_figures, start=1):
        add_figure(main_doc, display, f"Figure {number}")
    main_doc.save(MAIN_PATH)

    supp_doc = Document()
    configure_document(supp_doc, body_size=11, line_spacing=1.5)
    add_text(supp_doc, title, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(supp_doc, str(PROFILE_CONFIG.get("supplement_title", "ONLINE-ONLY SUPPLEMENTAL MATERIAL")), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(supp_doc, "[AUTHOR ACTION REQUIRED: confirm that all supplement citations, table labels, and figure labels match the final main manuscript]", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    supp_doc.add_page_break()
    add_events(supp_doc, supplement_events, include_displays=True, compact=True)
    supp_doc.save(SUPP_PATH)

    pdf_report = render_submission_pdfs()
    write_common_receipts(main_events, main_displays, supplement_events, running_title, evidence, pdf_report)
    return MAIN_PATH, SUPP_PATH


def main() -> None:
    main_path, supp_path = build()
    print(f"Wrote {main_path}")
    print(f"Wrote {supp_path}")
    print(f"Wrote section snapshots to {DRAFT_SECTION_ROOM}")
    print(f"Wrote source manifest to {MANIFEST_PATH}")
    print(f"Wrote QA report to {QA_REPORT_PATH}")
    print(f"Citation count: {len(CITATION_NUMBERS)}")


if __name__ == "__main__":
    main()
